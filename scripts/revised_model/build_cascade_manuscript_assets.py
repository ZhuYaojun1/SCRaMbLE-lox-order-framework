from __future__ import annotations

import argparse
import math
from pathlib import Path

import matplotlib as mpl
mpl.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
ANALYSIS = REPOSITORY_ROOT / "outputs" / "revised_model" / "cascade_analysis"
OUT = ANALYSIS / "manuscript_assets" / "rebuilt"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build the Section 3.8 cascade figures, source tables, and Word package."
    )
    parser.add_argument("--analysis-dir", type=Path, default=ANALYSIS)
    parser.add_argument("--output-dir", type=Path, default=OUT)
    return parser.parse_args()


COLORS = {
    "navy": "#315A7D",
    "blue": "#5F8FB8",
    "teal": "#4E9A8D",
    "orange": "#D68A45",
    "coral": "#C95D5D",
    "gold": "#D6B04C",
    "purple": "#8B78A8",
    "gray": "#A7ADB4",
    "light_gray": "#D7DADD",
    "dark": "#2B2B2B",
}


def configure_matplotlib() -> None:
    mpl.rcParams.update(
        {
            "font.family": "serif",
            "font.serif": ["Times New Roman", "Times", "DejaVu Serif"],
            "font.size": 12,
            "axes.labelsize": 12,
            "axes.titlesize": 12,
            "xtick.labelsize": 12,
            "ytick.labelsize": 12,
            "legend.fontsize": 12,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "axes.linewidth": 0.8,
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "svg.fonttype": "none",
        }
    )


def save_figure(fig: plt.Figure, stem: str) -> None:
    fig.savefig(OUT / f"{stem}.png", dpi=400, bbox_inches="tight", facecolor="white")
    fig.savefig(OUT / f"{stem}.pdf", bbox_inches="tight", facecolor="white")
    fig.savefig(OUT / f"{stem}.svg", bbox_inches="tight", facecolor="white")
    plt.close(fig)


def panel_label(ax: plt.Axes, label: str) -> None:
    ax.text(-0.14, 1.08, label, transform=ax.transAxes, fontsize=13, fontweight="bold", va="top")


def load_data() -> dict[str, pd.DataFrame]:
    files = {
        "sequence": "cascade_sequence_summary.csv",
        "mechanism": "cascade_mechanism_summary.csv",
        "amplification": "cascade_amplification_threshold_summary.csv",
        "event": "cascade_event_effect_summary.csv",
        "trajectory": "cascade_trajectory_effect_summary.csv",
        "hazard": "cascade_total_hazard_sample_summary.csv",
        "top": "cascade_top_probability_amplifications.csv",
    }
    return {key: pd.read_csv(ANALYSIS / filename) for key, filename in files.items()}


def figure_10(data: dict[str, pd.DataFrame]) -> None:
    sequence = data["sequence"]
    mechanism = data["mechanism"]
    amplification = data["amplification"]

    history_labels = {
        "deletion_only": "Deletion only",
        "inversion_only_history": "Inversion history",
        "earlier_deletion_with_nondeletion": "Earlier deletion +\nnon-deletion",
        "earlier_deletion_only": "Earlier deletion only",
        "mixed_inversion_duplication": "Mixed inversion–\nduplication",
        "duplication_only_history": "Duplication history",
    }
    history = sequence[sequence["scope"].eq("overall")].copy()
    history["label"] = history["history_class"].map(history_labels)
    history = history.sort_values("gate_failing_trajectories", ascending=True)
    history["percent"] = 100.0 * history["gate_failing_trajectories"] / history["gate_failing_trajectories"].sum()
    history[["history_class", "label", "gate_failing_trajectories", "percent"]].to_csv(
        OUT / "figure10_panel_a_source.csv", index=False
    )

    mechanism_labels = {
        "no_direct_effect_on_final_pair": "No direct effect",
        "proposal_space_renormalization": "Pair-space renormalization",
        "distance_shortening": "Distance shortening",
        "distance_lengthening": "Distance lengthening",
        "final_pair_activation": "New terminal pair",
        "lethal_span_activation": "Lethal-span activation",
        "essential_copy_buffering": "Copy-number buffering",
    }
    mechanism = mechanism.copy()
    mechanism["label"] = mechanism["cascade_mechanism"].map(mechanism_labels)
    mechanism.to_csv(OUT / "figure10_panel_b_source.csv", index=False)
    amplification.to_csv(OUT / "figure10_panel_c_source.csv", index=False)

    fig = plt.figure(figsize=(8.0, 8.0), constrained_layout=False)
    grid = fig.add_gridspec(2, 2, height_ratios=[1.05, 1.0], hspace=0.56, wspace=0.48)
    ax_a = fig.add_subplot(grid[0, 0])
    ax_b = fig.add_subplot(grid[0, 1])
    ax_c = fig.add_subplot(grid[1, :])

    bar_colors = [COLORS["gray"] if value in {"deletion_only", "earlier_deletion_only"} else COLORS["navy"] for value in history["history_class"]]
    ax_a.barh(history["label"], history["gate_failing_trajectories"] / 1000.0, color=bar_colors, height=0.68)
    for y, (_, row) in enumerate(history.iterrows()):
        ax_a.text(row["gate_failing_trajectories"] / 1000.0 + 3, y, f"{row['percent']:.1f}%", va="center", fontsize=11)
    ax_a.set_xlabel("Gate-failing trajectories (×10³)")
    ax_a.set_title("History composition", loc="left", fontweight="bold")
    ax_a.set_xlim(0, history["gate_failing_trajectories"].max() / 1000.0 * 1.25)
    ax_a.grid(axis="x", color="#E6E6E6", linewidth=0.7)
    ax_a.set_axisbelow(True)
    panel_label(ax_a, "A")

    categories = [
        "No direct effect",
        "Pair-space renormalization",
        "Distance shortening",
        "Distance lengthening",
        "New terminal pair",
        "Lethal-span activation",
        "Copy-number buffering",
    ]
    category_colors = {
        "No direct effect": COLORS["light_gray"],
        "Pair-space renormalization": COLORS["purple"],
        "Distance shortening": COLORS["teal"],
        "Distance lengthening": COLORS["gold"],
        "New terminal pair": COLORS["orange"],
        "Lethal-span activation": COLORS["coral"],
        "Copy-number buffering": COLORS["blue"],
    }
    event_order = ["inversion", "duplication"]
    event_labels = ["Inversion", "Duplication"]
    left = np.zeros(2)
    for category in categories:
        values = []
        for event_type in event_order:
            row = mechanism[(mechanism["precursor_event_type"].eq(event_type)) & (mechanism["label"].eq(category))]
            values.append(100.0 * float(row["event_fraction_within_type"].iloc[0]) if len(row) else 0.0)
        ax_b.barh(event_labels, values, left=left, color=category_colors[category], height=0.62, label=category)
        left += np.asarray(values)
    ax_b.set_xlim(0, 100)
    ax_b.set_xlabel("Precursor events (%)")
    ax_b.set_title("Immediate terminal-pair effect", loc="left", fontweight="bold")
    ax_b.grid(axis="x", color="#E6E6E6", linewidth=0.7)
    ax_b.set_axisbelow(True)
    handles, labels = ax_b.get_legend_handles_labels()
    panel_label(ax_b, "B")

    amp = amplification.copy()
    selected = pd.concat(
        [
            amp[(amp["precursor_event_type"].eq("inversion")) & amp["change_class"].isin(["at_least_2_fold", "at_least_10_fold", "at_least_100_fold"])],
            amp[(amp["precursor_event_type"].eq("duplication")) & amp["change_class"].eq("new_pair_activation")],
        ],
        ignore_index=True,
    )
    label_map = {
        "at_least_2_fold": "Inversion\n≥2-fold",
        "at_least_10_fold": "Inversion\n≥10-fold",
        "at_least_100_fold": "Inversion\n≥100-fold",
        "new_pair_activation": "Duplication\nnew pair",
    }
    selected["label"] = selected["change_class"].map(label_map)
    x = np.arange(len(selected))
    colors = [COLORS["teal"], COLORS["teal"], COLORS["teal"], COLORS["orange"]]
    ax_c.bar(x, selected["n_precursor_events"], color=colors, width=0.62)
    ax_c.set_yscale("log")
    ax_c.set_ylabel("Precursor events (log scale)")
    ax_c.set_xticks(x, selected["label"])
    ax_c.set_title("Large changes in eventual terminal-pair probability", loc="left", fontweight="bold")
    ax_c.grid(axis="y", which="major", color="#E6E6E6", linewidth=0.7)
    ax_c.set_axisbelow(True)
    for xpos, value in zip(x, selected["n_precursor_events"]):
        ax_c.text(xpos, value * 1.18, f"{int(value):,}", ha="center", va="bottom", fontsize=11)
    ax_c.set_ylim(4, 6500)
    panel_label(ax_c, "C")

    fig.legend(
        handles,
        labels,
        loc="lower center",
        bbox_to_anchor=(0.5, 0.01),
        ncol=3,
        fontsize=9.5,
        handlelength=1.2,
        columnspacing=1.0,
    )
    fig.subplots_adjust(left=0.14, right=0.97, top=0.93, bottom=0.20)
    save_figure(fig, "Figure_10_rearrangement_cascade")


def figure_11(data: dict[str, pd.DataFrame]) -> None:
    hazard = data["hazard"]
    detailed_path = ANALYSIS / "cascade_precursor_events.csv.xz"
    chunks = []
    usecols = ["model_key", "precursor_event_type", "total_hazard_sampled", "total_hazard_log2_ratio_nonzero"]
    for chunk in pd.read_csv(detailed_path, usecols=usecols, chunksize=50_000):
        sampled = chunk[chunk["total_hazard_sampled"].astype(str).str.lower().eq("true")].copy()
        chunks.append(sampled)
    detailed = pd.concat(chunks, ignore_index=True)
    detailed.to_csv(OUT / "figure11_source_events.csv", index=False)

    selected = hazard[~hazard["model_key"].eq("ALL")].copy()
    label_map = {
        ("linear_distance", "inversion"): "Linear\ninversion",
        ("partial_hic_fallback", "inversion"): "Partial Hi-C\ninversion",
        ("linear_distance", "duplication"): "Linear\nduplication",
        ("partial_hic_fallback", "duplication"): "Partial Hi-C\nduplication",
    }
    order = list(label_map)
    rows = []
    for model, event_type in order:
        row = selected[(selected["model_key"].eq(model)) & (selected["precursor_event_type"].eq(event_type))].iloc[0].to_dict()
        row["label"] = label_map[(model, event_type)]
        rows.append(row)
    plot = pd.DataFrame(rows)
    plot.to_csv(OUT / "figure11_panel_a_source.csv", index=False)

    fig, (ax_a, ax_b) = plt.subplots(1, 2, figsize=(8.0, 4.6), gridspec_kw={"width_ratios": [1.12, 1.0]})
    y = np.arange(len(plot))[::-1]
    increase = 100.0 * plot["hazard_increase_events"] / plot["n_sampled_events"]
    decrease = 100.0 * plot["hazard_decrease_events"] / plot["n_sampled_events"]
    unchanged = 100.0 * plot["hazard_unchanged_events"] / plot["n_sampled_events"]
    ax_a.barh(y, increase, color=COLORS["coral"], label="Increase")
    ax_a.barh(y, unchanged, left=increase, color=COLORS["gray"], label="Unchanged")
    ax_a.barh(y, decrease, left=increase + unchanged, color=COLORS["blue"], label="Decrease")
    ax_a.set_yticks(y, plot["label"])
    ax_a.set_xlabel("Sampled precursor states (%)")
    ax_a.set_xlim(0, 100)
    ax_a.set_title("Direction of one-step hazard change", loc="left", fontweight="bold")
    ax_a.legend(loc="upper center", bbox_to_anchor=(0.5, -0.22), ncol=3, fontsize=10, columnspacing=0.8)
    ax_a.grid(axis="x", color="#E6E6E6", linewidth=0.7)
    ax_a.set_axisbelow(True)
    panel_label(ax_a, "A")

    labels = []
    medians = []
    lows = []
    highs = []
    for model, event_type in order:
        values = pd.to_numeric(
            detailed[(detailed["model_key"].eq(model)) & (detailed["precursor_event_type"].eq(event_type))]["total_hazard_log2_ratio_nonzero"],
            errors="coerce",
        ).dropna()
        labels.append(label_map[(model, event_type)].replace("\n", " "))
        medians.append(float(values.median()))
        lows.append(float(values.quantile(0.025)))
        highs.append(float(values.quantile(0.975)))
    medians_arr = np.asarray(medians)
    ax_b.errorbar(
        medians_arr,
        y,
        xerr=np.vstack([medians_arr - np.asarray(lows), np.asarray(highs) - medians_arr]),
        fmt="o",
        color=COLORS["dark"],
        ecolor=COLORS["navy"],
        elinewidth=2,
        capsize=4,
        markersize=6,
    )
    ax_b.axvline(0, color=COLORS["gray"], linewidth=1, linestyle="--")
    ax_b.set_yticks(y, [])
    ax_b.set_xlabel("log₂(hazard after / hazard before)")
    ax_b.set_title("Median and central 95% of changes", loc="left", fontweight="bold")
    ax_b.grid(axis="x", color="#E6E6E6", linewidth=0.7)
    ax_b.set_axisbelow(True)
    panel_label(ax_b, "B")

    fig.subplots_adjust(left=0.19, right=0.98, top=0.87, bottom=0.26, wspace=0.34)
    save_figure(fig, "Figure_11_total_gate_failing_hazard")


def table_source(data: dict[str, pd.DataFrame]) -> pd.DataFrame:
    event = data["event"]
    trajectory = data["trajectory"]
    amplification = data["amplification"]
    hazard = data["hazard"]
    top = data["top"]
    overall_event = event[event["scope"].eq("overall")].set_index("precursor_event_type")
    overall_traj = trajectory[trajectory["model_key"].eq("ALL")].set_index("precursor_event_type")
    overall_hazard = hazard[hazard["model_key"].eq("ALL")].set_index("precursor_event_type")

    def pct(n: int, d: int) -> str:
        return f"{100.0 * n / d:.2f}%"

    rows = [
        ["Trajectory history", "Any preceding inversion or duplication", "196,246 / 438,708 trajectories", "44.73% of Gate-failing trajectories"],
        ["Event replay", "Inversion precursors", "241,208 events in 162,430 trajectories", "All events before the terminal deletion"],
        ["Event replay", "Duplication precursors", "86,890 events in 74,133 trajectories", "All events before the terminal deletion"],
        ["Exact terminal pair", "Hazard increased after inversion", "8,836 / 241,208 events", "3.66% (95% Wilson interval, 3.59–3.74%)"],
        ["Exact terminal pair", "Hazard increased after duplication", "3,270 / 86,890 events", "3.76% (95% Wilson interval, 3.64–3.89%)"],
        ["Trajectory level", "At least one inversion-associated increase", "8,597 / 162,430 trajectories", "5.29% (95% Wilson interval, 5.18–5.40%)"],
        ["Trajectory level", "At least one duplication-associated increase", "3,270 / 74,133 trajectories", "4.41% (95% Wilson interval, 4.27–4.56%)"],
        ["Finite probability amplification", "Inversion, ≥2-fold / ≥10-fold / ≥100-fold", "1,155 / 110 / 7 events", "Maximum finite increase, 248.06-fold"],
        ["Copy-specific pair creation", "Duplication activated the eventual terminal pair", "3,437 / 86,890 events", "3.96%"],
        ["Gate-span change", "Inversion activated a Gate-failing deletion span", "705 / 241,208 events", "0.29%"],
        ["All-pair one-step hazard sample", "Hazard increased after inversion", "1,833 / 4,400 states", "41.66%; 16.07% decreased and 42.27% were unchanged"],
        ["All-pair one-step hazard sample", "Hazard decreased after duplication", "4,239 / 4,385 states", "96.67%; 3.33% increased"],
    ]
    frame = pd.DataFrame(rows, columns=["Analysis level", "Rearrangement class or measure", "Count", "Result"])
    frame.to_csv(OUT / "Table_11_source.csv", index=False)
    return frame


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "bottom", "insideH", "insideV", "left", "right"]:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        if edge in {"top", "bottom"}:
            tag.set(qn("w:val"), "single")
            tag.set(qn("w:sz"), "10")
            tag.set(qn("w:color"), "000000")
        else:
            tag.set(qn("w:val"), "nil")
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "000000")
        tc_borders.append(bottom)


def format_run(run, size: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_caption(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(label)
    format_run(run, 10, bold=True)
    run = paragraph.add_run(text)
    format_run(run, 10)


def build_docx(table_frame: pd.DataFrame) -> Path:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    title = document.add_paragraph()
    title.style = normal
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Section 3.8 Figure and Table Package")
    format_run(run, 12, bold=True)
    title.paragraph_format.space_after = Pt(10)

    table_title = document.add_paragraph()
    table_title.style = normal
    table_title.paragraph_format.keep_with_next = True
    run = table_title.add_run("Table 11. ")
    format_run(run, 10, bold=True)
    run = table_title.add_run("State-replay summary of rearrangement-cascade effects preceding terminal Essentiality Gate failure.")
    format_run(run, 10)

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(1.25), Inches(2.15), Inches(1.55), Inches(1.95)]
    headers = list(table_frame.columns)
    for index, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = widths[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        format_run(run, 10, bold=True)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    for record in table_frame.to_dict("records"):
        row = table.add_row()
        for index, header in enumerate(headers):
            cell = row.cells[index]
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(record[header]))
            format_run(run, 10)
    set_table_borders(table)

    note = document.add_paragraph()
    note.style = normal
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(8)
    run = note.add_run("Note: ")
    format_run(run, 10, italic=True)
    run = note.add_run(
        "Exact terminal-pair analyses included every inversion or duplication preceding Gate failure. "
        "The all-pair hazard analysis used a deterministic scenario-stratified sample of up to 50 events per event type and scenario. "
        "Percentages are descriptive and are conditional on trajectories that later failed the strict Essentiality Gate."
    )
    format_run(run, 10)

    document.add_page_break()
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_with_next = True
    picture.add_run().add_picture(str(OUT / "Figure_10_rearrangement_cascade.png"), width=Inches(6.75))
    add_caption(
        document,
        "Figure 10. ",
        "Precursor rearrangements altered the accessibility of later Gate-failing deletion paths. "
        "(A) Composition of event histories among 438,708 Gate-failing trajectories. "
        "(B) Mutually exclusive classifications of 241,208 inversion and 86,890 duplication precursor events according to their immediate effect on the eventual terminal-deletion pair. "
        "(C) Inversion events producing finite increases of at least 2-, 10-, or 100-fold in the proposal probability of the eventual terminal pair; duplications that created a previously absent copy-specific terminal pair are shown separately."
    )

    document.add_page_break()
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_with_next = True
    picture.add_run().add_picture(str(OUT / "Figure_11_total_gate_failing_hazard.png"), width=Inches(6.75))
    add_caption(
        document,
        "Figure 11. ",
        "Precursor events produced event- and model-dependent changes in the total one-step Gate-failing deletion hazard. "
        "(A) Direction of hazard change after inversion or duplication under linear-distance sampling and partially Hi-C-informed sampling with distance-based fallback. "
        "(B) Median and central 95% of the descriptive distribution of log₂-transformed after-to-before hazard ratios. "
        "The scenario-stratified sample comprised 8,785 precursor states; intervals in panel B are distributional quantiles rather than confidence intervals."
    )

    output = OUT / "SCRaMbLE_section_3_8_Figure_Table_package.docx"
    document.save(output)
    return output


def main() -> None:
    global ANALYSIS, OUT
    args = parse_args()
    ANALYSIS = args.analysis_dir.resolve()
    OUT = args.output_dir.resolve()
    OUT.mkdir(parents=True, exist_ok=True)
    configure_matplotlib()
    data = load_data()
    figure_10(data)
    figure_11(data)
    table_frame = table_source(data)
    output = build_docx(table_frame)
    print(output)


if __name__ == "__main__":
    main()
